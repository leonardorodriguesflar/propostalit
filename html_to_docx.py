from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup
import os

# Caminhos dos arquivos
html_path = "Ominicom/index.HTML"
img_path = "Ominicom/Lit - Banners_Prancheta 1 (1).jpg"
docx_path = "Ominicom/Apresentacao.docx"

# Lê o HTML
with open(html_path, "r", encoding="utf-8") as f:
    soup = BeautifulSoup(f, "html.parser")
    texto = soup.get_text(separator="\n", strip=True)

# Cria documento Word
doc = Document()
doc.add_paragraph(texto)

# Adiciona imagem (ajuste o tamanho se necessário)
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(2))

# Salva o arquivo docx
doc.save(docx_path)
print(f"Arquivo DOCX gerado em: {docx_path}") 